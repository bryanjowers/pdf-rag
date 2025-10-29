#!/usr/bin/env python3
"""
docx.py - Word document processing handler using Docling

Processes DOCX files using Docling's document converter with python-docx fallback.
Preserves sections, tables, and footnotes in markdown output.
"""

import json
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

from docling.document_converter import DocumentConverter


def process_docx(
    docx_path: Path,
    output_dir: Path,
    config: Dict,
    batch_id: str
) -> Dict:
    """
    Process DOCX file using Docling with fallback to python-docx.

    Args:
        docx_path: Path to input DOCX file
        output_dir: Base output directory
        config: Configuration dictionary
        batch_id: Unique batch identifier

    Returns:
        Processing result dictionary with metadata:
        {
            "success": bool,
            "processor": "docling" | "python-docx" | None,
            "markdown_path": Path | None,
            "jsonl_path": Path | None,
            "processing_duration_ms": int,
            "char_count": int,
            "estimated_tokens": int,
            "chunk_count": int,
            "warnings": list[str],
            "error": str | None
        }

    Raises:
        FileNotFoundError: If DOCX doesn't exist
    """
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX not found: {docx_path}")

    start_time = time.time()
    warnings = []

    # Prepare output paths
    markdown_dir = output_dir / "markdown"
    jsonl_dir = output_dir / "jsonl"
    markdown_dir.mkdir(parents=True, exist_ok=True)
    jsonl_dir.mkdir(parents=True, exist_ok=True)

    stem = docx_path.stem
    markdown_path = markdown_dir / f"{stem}.md"
    jsonl_path = jsonl_dir / f"{stem}.jsonl"

    # Try Docling first
    try:
        print(f"   🔄 Converting with Docling: {docx_path.name}")
        result = _process_with_docling(docx_path, markdown_path, jsonl_path, config, batch_id)

        if result["success"]:
            duration_ms = int((time.time() - start_time) * 1000)
            result["processing_duration_ms"] = duration_ms
            print(f"   ✅ Docling processing complete")
            print(f"      Output: {result['char_count']:,} chars (~{result['estimated_tokens']:,} tokens)")
            print(f"      Duration: {duration_ms/1000:.1f}s")
            print(f"      Chunks: {result['chunk_count']}")
            return result

    except Exception as e:
        warnings.append(f"Docling failed: {str(e)}")
        print(f"   ⚠️  Docling failed: {e}")

    # Fallback to python-docx
    print(f"   🔄 Attempting fallback to python-docx...")
    try:
        result = _process_with_python_docx(docx_path, markdown_path, jsonl_path, config, batch_id)

        if result["success"]:
            duration_ms = int((time.time() - start_time) * 1000)
            result["processing_duration_ms"] = duration_ms
            result["warnings"] = warnings
            print(f"   ✅ python-docx processing complete")
            print(f"      Output: {result['char_count']:,} chars (~{result['estimated_tokens']:,} tokens)")
            print(f"      Duration: {duration_ms/1000:.1f}s")
            print(f"      Chunks: {result['chunk_count']}")
            return result

    except Exception as e:
        duration_ms = int((time.time() - start_time) * 1000)
        error_msg = f"Both Docling and python-docx failed: {str(e)}"
        print(f"   ❌ {error_msg}")

        return {
            "success": False,
            "processor": None,
            "markdown_path": None,
            "jsonl_path": None,
            "processing_duration_ms": duration_ms,
            "char_count": 0,
            "estimated_tokens": 0,
            "chunk_count": 0,
            "warnings": warnings,
            "error": error_msg
        }


def _process_with_docling(
    docx_path: Path,
    markdown_path: Path,
    jsonl_path: Path,
    config: Dict,
    batch_id: str
) -> Dict:
    """
    Process DOCX using Docling converter.

    Args:
        docx_path: Input DOCX path
        markdown_path: Output markdown path
        jsonl_path: Output JSONL path
        config: Configuration dictionary
        batch_id: Batch identifier

    Returns:
        Processing result dictionary

    Raises:
        Exception: If conversion fails
    """
    # Initialize Docling converter
    converter = DocumentConverter()

    # Convert DOCX to document
    result = converter.convert(str(docx_path))

    # Extract markdown
    markdown_content = result.document.export_to_markdown()
    char_count = len(markdown_content)

    # Try to get page count from Docling metadata
    page_count = None
    if hasattr(result.document, 'page_count'):
        page_count = result.document.page_count
    elif hasattr(result.document, 'pages'):
        page_count = len(result.document.pages)
    else:
        # Estimate: ~300 words per page (typical for legal docs with 1" margins)
        word_count = len(markdown_content.split())
        page_count = max(1, round(word_count / 300))

    # Write markdown output
    markdown_path.write_text(markdown_content, encoding="utf-8")

    # Convert to JSONL chunks
    chunks = _convert_to_jsonl(
        markdown_content,
        docx_path,
        config,
        batch_id,
        processor="docling",
        file_type="docx"
    )

    # Write JSONL output
    with jsonl_path.open("w", encoding="utf-8") as f:
        for chunk in chunks:
            f.write(json.dumps(chunk, ensure_ascii=False) + "\n")

    return {
        "success": True,
        "processor": "docling",
        "markdown_path": markdown_path,
        "jsonl_path": jsonl_path,
        "processing_duration_ms": None,  # Set by caller
        "page_count": page_count,  # Estimated from word count if not available
        "char_count": char_count,
        "estimated_tokens": len(markdown_content.split()),
        "chunk_count": len(chunks),
        "warnings": [],
        "error": None
    }


def _process_with_python_docx(
    docx_path: Path,
    markdown_path: Path,
    jsonl_path: Path,
    config: Dict,
    batch_id: str
) -> Dict:
    """
    Fallback processing using python-docx library.

    Args:
        docx_path: Input DOCX path
        markdown_path: Output markdown path
        jsonl_path: Output JSONL path
        config: Configuration dictionary
        batch_id: Batch identifier

    Returns:
        Processing result dictionary

    Raises:
        Exception: If conversion fails
    """
    from docx import Document

    # Load document
    doc = Document(str(docx_path))

    # Extract text with basic formatting
    markdown_lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect headings (basic heuristic)
        if para.style.name.startswith('Heading'):
            level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
            markdown_lines.append(f"{'#' * level} {text}")
        else:
            markdown_lines.append(text)

        markdown_lines.append("")  # Blank line

    # Extract tables
    for table in doc.tables:
        markdown_lines.append("")  # Blank line before table

        # Convert table to markdown
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            markdown_lines.append("| " + " | ".join(cells) + " |")

            # Add header separator after first row
            if i == 0:
                markdown_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")

        markdown_lines.append("")  # Blank line after table

    markdown_content = "\n".join(markdown_lines)
    char_count = len(markdown_content)

    # Estimate page count from word count (~300 words per page for legal docs)
    word_count = len(markdown_content.split())
    page_count = max(1, round(word_count / 300))

    # Write markdown output
    markdown_path.write_text(markdown_content, encoding="utf-8")

    # Convert to JSONL chunks
    chunks = _convert_to_jsonl(
        markdown_content,
        docx_path,
        config,
        batch_id,
        processor="python-docx",
        file_type="docx"
    )

    # Write JSONL output
    with jsonl_path.open("w", encoding="utf-8") as f:
        for chunk in chunks:
            f.write(json.dumps(chunk, ensure_ascii=False) + "\n")

    return {
        "success": True,
        "processor": "python-docx",
        "markdown_path": markdown_path,
        "jsonl_path": jsonl_path,
        "processing_duration_ms": None,  # Set by caller
        "page_count": page_count,  # Estimated from word count
        "char_count": char_count,
        "estimated_tokens": len(markdown_content.split()),
        "chunk_count": len(chunks),
        "warnings": [],
        "error": None
    }


def _convert_to_jsonl(
    markdown_text: str,
    source_path: Path,
    config: Dict,
    batch_id: str,
    processor: str,
    file_type: str
) -> list[Dict]:
    """
    Convert markdown text to JSONL chunks following unified schema.

    Args:
        markdown_text: Markdown content to chunk
        source_path: Original source file path
        config: Configuration dictionary
        batch_id: Batch identifier
        processor: Name of processor used
        file_type: Type of source file (docx, pdf, etc.)

    Returns:
        List of JSONL record dictionaries
    """
    import sys
    sys.path.insert(0, str(Path(__file__).parent.parent))
    from utils_classify import compute_file_hash, get_mime_type

    # Compute source file hash
    file_hash = compute_file_hash(source_path)
    mime_type = get_mime_type(source_path)

    # Generate doc_id from hash (first 16 chars)
    doc_id = file_hash[:16]

    # Simple chunking by paragraphs
    paragraphs = [p.strip() for p in markdown_text.split('\n\n') if p.strip()]

    # Get chunking config
    chunking_config = config.get("chunking", {})
    token_target = chunking_config.get("token_target", 1400)
    token_max = chunking_config.get("token_max", 2000)

    # Combine paragraphs into chunks within token limits
    chunks = []
    current_chunk = []
    current_tokens = 0

    for para in paragraphs:
        para_tokens = len(para.split())

        # If adding this paragraph exceeds max, finalize current chunk
        if current_tokens + para_tokens > token_max and current_chunk:
            chunks.append("\n\n".join(current_chunk))
            current_chunk = [para]
            current_tokens = para_tokens
        else:
            current_chunk.append(para)
            current_tokens += para_tokens

            # If we've reached target size, finalize chunk
            if current_tokens >= token_target:
                chunks.append("\n\n".join(current_chunk))
                current_chunk = []
                current_tokens = 0

    # Add remaining chunk
    if current_chunk:
        chunks.append("\n\n".join(current_chunk))

    # Convert to JSONL records
    schema_version = config.get("schema", {}).get("version", "2.2.0")
    processed_at = datetime.utcnow().isoformat() + "Z"

    jsonl_records = []
    for idx, chunk_text in enumerate(chunks):
        chunk_tokens = len(chunk_text.split())

        record = {
            "id": f"{doc_id}_{idx:04d}",
            "doc_id": doc_id,
            "chunk_index": idx,
            "text": chunk_text,
            "attrs": {
                "page_span": None,
                "sections": [],
                "table": "| " in chunk_text and chunk_text.count("|") >= 3,  # Simple table detection
                "token_count": chunk_tokens
            },
            "source": {
                "file_path": str(source_path.resolve()),
                "file_name": source_path.name,
                "file_type": file_type,
                "mime_type": mime_type
            },
            "metadata": {
                "schema_version": schema_version,
                "file_type": file_type,
                "mime_type": mime_type,
                "hash_input_sha256": file_hash,
                "processor": processor,
                "processor_version": "docling-latest" if processor == "docling" else "python-docx-1.0",
                "processed_at": processed_at,
                "batch_id": batch_id,
                "processing_duration_ms": None,
                "confidence_score": 1.0,
                "warnings": []
            }
        }

        jsonl_records.append(record)

    return jsonl_records
